import openai
from docx import Document
import os
import PyPDF2

openai.api_key = ''

def generate_learning_plan_doc(curriculum_text, standard_text, weeks, sessions, hours):
    """
    Generate a learning plan DOCX file using OpenAI + python-docx (real table).
    """
    # Step 1: Ask OpenAI for structured data
    prompt = f"""
    You are tasked with generating a structured learning plan.
    
    Curriculum: {curriculum_text}
    Occupational Standard: {standard_text}
    
    Duration: {weeks} weeks
    Sessions per week: {sessions}
    Hours per session: {hours}
    
    Return the output as rows in a table with columns:
    Week | Session | Topic | Session Objectives | Trainer Activities | Trainee Activities | Learning Aids & References |Assessment.
    On session objectives use verbs from Bloom's taxonomy and there should be at least two objectives and not more than 4
    On Trainer and trainee activities, also use the verbs from Bloom's taxonomy
    On Learning aids, use appropriate aids including use of projectors, IDE's, and on references, give at least two references using the api format
    Generate for all the 12 weeks as provided above with each week having 2 sessions
    """
    
    response = openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2500
    )

    plan_text = response.choices[0].message.content.strip()
    
    print(plan_text)
    # Step 2: Parse response into rows
    # rows = []
    # for line in plan_text.split("\n"):
    #     parts = [p.strip() for p in line.split(",")]
    #     if len(parts) >= 5:  # Ensure 5 columns
    #         rows.append(parts[:5])

    # # Step 3: Create Word document
    # doc = Document()
    # doc.add_heading("Generated Learning Plan", level=1)

    # # Create table
    # table = doc.add_table(rows=1, cols=8)
    # table.style = "Table Grid"

    # # Add headers
    # hdr_cells = table.rows[0].cells
    # headers = ["Week", "Session", "Topic", "Session Objectives", "Trainer Activities", "Trainee Activities", "Learning Aids & References", "Assessment"]
    # for i, h in enumerate(headers):
    #     hdr_cells[i].text = h

    # # Add rows
    # for row in rows:
    #     cells = table.add_row().cells
    #     for i, val in enumerate(row):
    #         cells[i].text = val

    # # Step 4: Save file
    # os.makedirs("media", exist_ok=True)
    # output_path = os.path.join("media", "learning_plan.docx")
    # doc.save(output_path)

    return plan_text.encode('utf-8')  # Return bytes of the document

# curriculum_path = 'TrainerProject/TrainerApp/mpesa_utils/aicr.pdf'
# standard_path = 'TrainerProject/TrainerApp/mpesa_utils/ai_os.pdf'

# curriculum = ""
# standard = ""

# try:
#     # Open the PDF file in binary mode ('rb')
#     with open(curriculum_path, 'rb') as file:
#         # Create a PDF reader object
#         reader = PyPDF2.PdfReader(file)

#         # Get the number of pages in the PDF
#         num_pages = len(reader.pages)

#         # Iterate through all the pages and extract the text
#         for page_num in range(num_pages):
#             page = reader.pages[page_num]
#             curriculum += page.extract_text()
            
#         curriculum = curriculum.encode("ascii", "ignore").decode("ascii")
#     # Now you can use the 'curriculum' variable which contains all the text from the PDF
#     # print(curriculum)
    
    

# except FileNotFoundError:
#     print(f"Error: The file at {curriculum_path} was not found.")
# except Exception as e:
#     print(f"An error occurred: {e}")
    
# try:
#     # Open the PDF file in binary mode ('rb')
#     with open(standard_path, 'rb') as file:
#         # Create a PDF reader object
#         reader = PyPDF2.PdfReader(file)

#         # Get the number of pages in the PDF
#         num_pages = len(reader.pages)

#         # Iterate through all the pages and extract the text
#         for page_num in range(num_pages):
#             page = reader.pages[page_num]
#             standard += page.extract_text()
            
#         standard = standard.encode("ascii", "ignore").decode("ascii")
#     # Now you can use the 'standard' variable which contains all the text from the PDF
#     # print(standard)
    
# except FileNotFoundError:
#     print(f"Error: The file at {standard_path} was not found.")
# except Exception as e:
#     print(f"An error occurred: {e}")

# # curriculum = open('TrainerProject/TrainerApp/mpesa_utils/ai_os.pdf', 'r').read()
# # standard = open('TrainerProject/TrainerApp/mpesa_utils/aicr.pdf', 'r').read()
# print(generate_learning_plan_doc(curriculum, standard, 4, 2, 3))
